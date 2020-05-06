# -*-coding:gbk*-
from os import listdir
import docx
# from docx.enum.table import WD_TABLE_ALIGNMENT  # �������Ǳ����뷽ʽ����ʱû��
# from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd


class excelToWord():
    def __init__(self):
        # �޸ĵı��λ��
        self.INDEX = 13
        # �ļ����λ��
        self.readfile_path = r'./0_example_doc/'
        self.excel_source = r'./1_excel_source/'
        # ���г���
        self.write_doc()
        pass

    def read_excel(self,doc):  # ��ȡexcel�ļ�
        df=pd.read_excel(self.excel_source+doc+'.xlsx',header=None)  # ��ȡ����������&������
        return df

    def write_doc(self):
        xlslist = listdir(self.excel_source)  # excel_source�ļ����������ļ�
        for myxls in xlslist:
            docname=myxls[:-5]  # ȥ����׺������ļ���
            df = self.read_excel(docname)  # ��ȡexcel�ļ�
            newdocx = docx.Document(self.readfile_path+'example.docx')
            # newdocx.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = newdocx.tables  # ��ȡdoc�е����б��
            for index,oTable in enumerate(table):
                if index!=self.INDEX:  # ֻ�޸ĵ�14����񣬶�Ӧӯ������
                    continue
                rows_num = len(oTable.rows)  # �������
                columns_num = len(oTable.columns)  # �������
                # ѭ�����У��޸ı��ֵ
                # �����ϣ���ı��������У����Ըı�ѭ����ʼ��
                for i in range(0,rows_num):  # ��ѭ��
                    for j in range(0,columns_num):  # ��ѭ��
                        cell=df.iloc[i,j]
                        if pd.isnull(cell)==True:  # ���Ϊ��ֵ�������ֵ
                            oTable.cell(i,j).text=''
                        else:
                            oTable.cell(i,j).text=str(cell)# д��cell
                        # ���������뷽ʽ
                        # oTable.cell(i,j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            newdocx.save('./2_new_doc/' + str(docname)+'.docx') # ����doc�����ļ�����
            print('save '+str(myxls))

if __name__ == '__main__':
    excelToWord()